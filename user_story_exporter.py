import streamlit as st
from docx import Document
from io import BytesIO
import pandas as pd

st.set_page_config(page_title="User Stories Generator", layout="wide")

st.title("📘 User Stories Generator met Tabel")

st.write("Voer hier eenvoudig meerdere user stories in via een tabel en exporteer naar Word.")

# Begin met lege tabel als er nog niks is
if "stories" not in st.session_state:
    st.session_state["stories"] = pd.DataFrame(
        [{"ID": "US-001", "Als": "", "Wil ik": "", "Zodat": "", "Kleur": "#FFFFFF"}]
    )

# Data editor voor invoer in tabel
st.session_state["stories"] = st.data_editor(
    st.session_state["stories"],
    num_rows="dynamic",
    use_container_width=True,
    key="stories_editor",
    hide_index=True,
    column_config={
        "Kleur": st.column_config.ColorColumn("Kleur", help="Kies een kleur voor deze user story"),
        "Als": st.column_config.TextColumn("Als"),
        "Wil ik": st.column_config.TextColumn("Wil ik"),
        "Zodat": st.column_config.TextColumn("Zodat"),
    },
)

# Knop om ID's automatisch te genereren
if st.button("🔢 Genereer ID’s automatisch"):
    st.session_state["stories"]["ID"] = [
        f"US-{i:03d}" for i in range(1, len(st.session_state["stories"]) + 1)
    ]

# Voorbeeldweergave
st.subheader("📋 Voorbeeldweergave")
for _, story in st.session_state["stories"].iterrows():
    st.markdown(
        f"""
        <div style="background-color:{story['Kleur']}; padding:12px; border-radius:10px; margin-bottom:12px;">
        <b>{story['ID']}</b><br>
        - Als {story['Als']}<br>
        - Wil ik {story['Wil ik']}<br>
        - Zodat {story['Zodat']}
        </div>
        """,
        unsafe_allow_html=True,
    )

# Export naar Word
if st.button("💾 Exporteer naar Word"):
    doc = Document()
    doc.add_heading("User Stories", level=1)

    for _, story in st.session_state["stories"].iterrows():
        doc.add_heading(story["ID"], level=2)
        doc.add_paragraph(f"Als {story['Als']}")
        doc.add_paragraph(f"Wil ik {story['Wil ik']}")
        doc.add_paragraph(f"Zodat {story['Zodat']}")
        doc.add_paragraph("---")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Download Word-document",
        data=buffer,
        file_name="user_stories.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
